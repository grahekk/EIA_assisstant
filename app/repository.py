from app import db
import geopandas as gpd, fiona
from shapely.geometry import Point
import pyproj
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
import markdown2
import pypandoc

import matplotlib.pyplot as plt
import cartopy.crs as ccrs


# fiona.drvsupport.supported_drivers['kml'] = 'rw' 
# fiona.drvsupport.supported_drivers['KML'] = 'rw' 
# fiona.drvsupport.supported_drivers['LIBKML'] = 'rw'
# fiona.drvsupport.supported_drivers['libkml'] = 'rw'
# fiona.drvsupport.supported_drivers['GPKG'] = 'rw'


def transform_coordinates(lat, lon, from_epsg=4326, to_epsg=3765):
    transformer = pyproj.Transformer.from_crs(from_epsg, to_epsg, always_xy=True)
    x, y = transformer.transform(lon, lat)
    return x, y

def check_intersection(lat, lon):
    x, y = transform_coordinates(lat, lon)

    geo_file = "3_parcijalni_ispit\pop.shp"
    geo_file = "pop.shp"
    # Read the CSV file into a GeoDataFrame
    gdf = gpd.read_file(geo_file)
    gdf = gdf.set_crs(3765, allow_override=True)

    # Create a Point geometry for the given lat, lon
    point = Point(x, y)

    # Check for intersection
    intersection = gdf.geometry.contains(point)

    # If there is an intersection, extract information
    if intersection.any():
        intersected_data = gdf[intersection]
        sitecode = intersected_data["sitecode"].iloc[0]
        print(sitecode)
        return sitecode
    else:
        return None


def natura_impact_assessment(lat, lon, project_title, project_type):
#     bird_data = "birds.csv"
#     bird_data = query_points_within_polygon(lat, lon)
    text = f"""Project named '{project_title}' is situated at the location {lat}, {lon}.
           As a {project_type} development project, it is crucial to consider the potential environmental impacts.
           
           The project could have significant effects on the local ecosystem, including wildlife and bird populations. 
           Some of the potential impacts are:
            
            1. Habitat Destruction: The development may lead to the destruction of natural habitats, affecting the local flora and fauna.
            
            2. Disruption of Wildlife Migration: If the project is located along migration routes, it could disrupt the natural migration patterns of birds and other animals.
            
            3. Noise Pollution: Construction activities can contribute to noise pollution, which may disturb bird species that rely on acoustic signals for communication and navigation.
            
            4. Air and Water Pollution: Depending on the nature of the project, there might be emissions and runoff that can contaminate the air and water, impacting both terrestrial and aquatic bird species.
            
            5. Collision Risks: Tall structures such as towers or wind turbines pose a collision risk for birds. It's essential to assess and mitigate these risks to protect bird populations.
            
            6. Introduction of Invasive Species: Construction activities may introduce invasive species that could outcompete or prey on local bird species.
            
            It's crucial to refer to available environmental impact assessments and conduct thorough studies to minimize negative consequences. The following bird species have been identified in the wider area of the project location based on available data in:

            Consideration of these factors is vital for sustainable development, ensuring that the project minimizes its ecological footprint and preserves biodiversity."""
            
    return text

def natura_description(project_title, site_code, site_name, distance, intersection: bool):
    if intersection == True:
        text = f"""Location of project {project_title} is situated at the location of Natura2000 protected areas:
                    The following areas are touched:
                    {site_code} {site_name}
                    The following areas are close:
                    {site_code} {site_name} which is distanced {distance}
                    ."""
                
    elif intersection == False:
        text = f"""Location of {project_title} is not situated on the locations of Natura2000 protected areas
                    The name of the area/s is:
                    {site_name} which is distanced {distance}km
                    ."""                
    return text


def apply_table_style(table):
    # Apply a built-in table style
    table.style = 'Table Grid'

    # Customize the style further if needed
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Customize font size if needed

    # Remove internal borders to improve style
    for row in table.rows:
        for cell in row.cells:
            for tc in cell._element.xpath('.//w:tcBorders'):
                tc.attrib.clear()
            

def create_report_docx(proj_name, text, table_data, image_path, output_path, merge_column=1):
    # Create a new Word document
    doc = Document()

    # Add a title with the project name
    doc.add_heading(proj_name, level=0)

    # Add the text content to the document
    heading = "POP - birds"
    doc.add_heading(heading, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(text)

    # add picture
    image_description = f"Figure: location of project {proj_name} in correlation with Natura2000 ecological framework"
    doc.add_picture(image_path, width=Inches(6))
    doc.add_paragraph(image_description).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create a footnote reference
    footnote_text = "Source: https://www.haop.hr/hr/tematska-podrucja/odrzivo-koristenje-prirodnih-dobara-i-ekoloska-mreza/ekoloska-mreza"
    footnote_reference = doc.add_footnote(footnote_text)
    # Add a footnote reference in the document
    footnote_reference = doc.add_footnote(footnote_reference)

    # Get the footnote reference number
    reference_number = footnote_reference.reference_id

    # Create a separate paragraph for the footnote text
    footnote_text = doc.add_paragraph(footnote_reference)

    # Set the font size of the footnote text
    for run in footnote_text.runs:
        font = run.font
        font.size = Pt(10)


    heading = "Table"
    doc.add_page_break()
    doc.add_heading(heading, level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a table to the document
    table = doc.add_table(rows=1, cols=len(table_data[0]))
    table.autofit = True

    # Add header row to the table
    header_row = table.rows[0].cells
    for col, header in enumerate(table_data[0]):
        header_row[col].text = str(header)

    # Add data rows to the table
    for row_data in table_data[1:]:
        row_cells = table.add_row().cells
        for col, value in enumerate(row_data):
            row_cells[col].text = str(value)

    merge_cells_in_column(table, column_index=1)
    
    # Apply a built-in table style
    apply_table_style(table)

    # Save the document to the specified output path
    doc.save(output_path)

def merge_cells_in_column(table, column_index):
    column_values = [cell.text for cell in table.columns[column_index].cells]
    merged_cells = set()

    for i, value in enumerate(column_values):
        if value not in merged_cells:
            # Find all cells with the same value in the column
            cells_to_merge = [cell for cell in table.columns[column_index].cells[i:] if cell.text == value]

            # Merge the cells
            if len(cells_to_merge) > 1:
                # Delete the content of all cells except the last one
                for cell in cells_to_merge[:-1]:
                    cell.text = ''
                
                # Update the set of merged cells
                merged_cells.add(value)


def generate_md_document(title, heading, paragraph, table, table_name, image, image_description, source):
    md_content = f"# {title}\n\n"

    # Heading
    md_content += f"## {heading}\n\n"

    # Paragraph
    md_content += f"{paragraph}\n\n"

    # Table
    md_content += f"Table: {table_name} \n"
    md_content += "| Header 1 | Header 2 |\n"
    md_content += "| -------- | -------- |\n"
    for row in table:
        md_content += f"| {row[0]} | {row[1]} |\n"
    md_content += "\n"

    # Image
    md_content += f"![{image_description}]({image})\n\n"

    # Source as footnote
    md_content += f"Source: {source}[^1]\n\n"

    # Footnote
    md_content += "[^1]: Source details go here\n"

    return md_content


def create_report(title, paragraph, table, image, output):
    chapter = "Natura"
    table_name = "Table 1: species on location of project\n"
    image_description = "Image 1: map of location of project\n"
    source = "https://www.haop.hr/hr/tematska-podrucja/odrzivo-koristenje-prirodnih-dobara-i-ekoloska-mreza/ekoloska-mreza\n"

    md_content = generate_md_document(title, chapter, paragraph, table, table_name, image, image_description, source)

    md_path = "output.md"
    with open(md_path, "w", encoding="utf-8") as md_file:
        md_file.write(md_content)
        
    md_to_docx(md_path, output)


def md_to_docx(md_file, docx_file):
    # Read the Markdown content from the file
    with open(md_file, 'r', encoding='utf-8') as md_file:
        md_content = md_file.read()

    # Convert Markdown to HTML, then html to docx
    html_content = markdown2.markdown(md_content, extras=['tables', 'footnotes'])
    with open("temporary_html_report.html", 'w', encoding='utf-8') as html_file:
        html_file.write(html_content)

    pypandoc.convert_text(html_content, to = "docx", format="html", outputfile=docx_file)


def report_knit(file_to_knit, output, format = 'docx'):
    pypandoc.convert_file(file_to_knit, to = format, outputfile=output)


if __name__ == "__main__":

    table_data = [
    ["ID", "Name", "Value"],
    [1, "Item 1", 10],
    [2, "Item 2", 20],
    [3, "Item 3", 30]
]
    