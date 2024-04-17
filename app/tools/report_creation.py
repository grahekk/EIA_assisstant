from app import db
import geopandas as gpd, fiona
from shapely.geometry import Point
import pyproj
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
import markdown2
import pypandoc

import matplotlib.pyplot as plt


# fiona.drvsupport.supported_drivers['kml'] = 'rw' 
# fiona.drvsupport.supported_drivers['KML'] = 'rw' 
# fiona.drvsupport.supported_drivers['LIBKML'] = 'rw'
# fiona.drvsupport.supported_drivers['libkml'] = 'rw'
# fiona.drvsupport.supported_drivers['GPKG'] = 'rw'


def transform_coordinates(lat, lon, from_epsg=4326, to_epsg=3765):
    transformer = pyproj.Transformer.from_crs(from_epsg, to_epsg, always_xy=True)
    x, y = transformer.transform(lon, lat)
    return x, y

def check_intersection(lat, lon):
    x, y = transform_coordinates(lat, lon)

    geo_file = "3_parcijalni_ispit\pop.shp"
    geo_file = "pop.shp"
    # Read the CSV file into a GeoDataFrame
    gdf = gpd.read_file(geo_file)
    gdf = gdf.set_crs(3765, allow_override=True)

    # Create a Point geometry for the given lat, lon
    point = Point(x, y)

    # Check for intersection
    intersection = gdf.geometry.contains(point)

    # If there is an intersection, extract information
    if intersection.any():
        intersected_data = gdf[intersection]
        sitecode = intersected_data["sitecode"].iloc[0]
        print(sitecode)
        return sitecode
    else:
        return None


def natura_impact_assessment(lat, lon, project_title, project_type):
#     bird_data = "birds.csv"
#     bird_data = query_points_within_polygon(lat, lon)
    text = f"""Predmetni zahvat za koji se provodi postupak PUO/OPP, '{project_title}' se nalazi na koordinatama {round(lat, 3)}, {round(lon,3)}.
           Budući da je to projekt tipa {project_type}, i podliježe procjeni prema Pravilniku o Procjeni utjecaja na okoliš (NN, 2019) bitno je razmotriti potencijalne utjecaje na okoliš.
           
           Projekt bi mogao imati utjecaja na neke sastavnice okoliša, tipa na lokalni ekosustav, populacije divljači i ptica.
            Neki od potencijalnih utjecaja su:
            
             1. Uništavanje staništa: Razvoj može dovesti do uništenja prirodnih staništa, utječući na lokalnu floru i faunu.
            
             2. Ometanje migracije divljih životinja: Ako se projekt nalazi duž migracijskih ruta, mogao bi poremetiti prirodne obrasce migracije ptica i drugih životinja.
            
             3. Zagađenje bukom: Građevinske aktivnosti mogu pridonijeti zagađenju bukom, što može uznemiriti vrste ptica koje se oslanjaju na akustične signale za komunikaciju i navigaciju.
            
             4. Onečišćenje zraka i vode: Ovisno o prirodi projekta, može doći do emisija i otjecanja koji mogu zagaditi zrak i vodu, utječući i na kopnene i na vodene vrste ptica.
            
             5. Rizici od sudara: Visoke strukture poput tornjeva ili vjetroturbina predstavljaju rizik od sudara za ptice. Neophodno je procijeniti i ublažiti ove rizike kako bi se zaštitile populacije ptica.
            
             6. Uvođenje invazivnih vrsta: Građevinske aktivnosti mogu uvesti invazivne vrste, ponajviše flore koje bi mogle ući u kompeticiju s autohtonom florom.
            
             Ključno je donijeti odgovarajuće mjere zaštite okoliša te provesti temeljita istraživanja kako bi se negativne posljedice svele na minimum. 

             Razmatranje svih čimbenika ključno je za održivi razvoj, time se osigurava da projekt minimalizira svoj okolišni otisak i očuva biološku raznolikost."""
            
    return text

def natura_description(project_title, site_code, site_name, distance, intersection: bool):
    if intersection == True:
        text = f"""Projekt {project_title} se nalazi na području ili pored Natura2000 zaštićenih područja:
                    {site_code} {site_name}."""
                
    elif intersection == False:
        text = f"""Location of {project_title} is not situated on the locations of Natura2000 protected areas
                    The name of the area/s is:
                    {site_name} which is distanced {distance}km
                    ."""                
    return text


def apply_table_style(table):
    # Apply a built-in table style
    table.style = 'Table Grid'

    # Customize the style further if needed
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Customize font size if needed

    # Remove internal borders to improve style
    for row in table.rows:
        for cell in row.cells:
            for tc in cell._element.xpath('.//w:tcBorders'):
                tc.attrib.clear()
            

def create_report_docx(proj_name, text, table_data, image_path, output_path, merge_column=1):
    # Create a new Word document
    doc = Document()

    # Add a title with the project name
    doc.add_heading(proj_name, level=0)

    # Add the text content to the document
    heading = "POP - birds"
    doc.add_heading(heading, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(text)

    # add picture
    image_description = f"Figure: location of project {proj_name} in correlation with Natura2000 ecological framework"
    doc.add_picture(image_path, width=Inches(6))
    doc.add_paragraph(image_description).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create a footnote reference
    footnote_text = "Source: https://www.haop.hr/hr/tematska-podrucja/odrzivo-koristenje-prirodnih-dobara-i-ekoloska-mreza/ekoloska-mreza"
    footnote_reference = doc.add_footnote(footnote_text)
    # Add a footnote reference in the document
    footnote_reference = doc.add_footnote(footnote_reference)

    # Get the footnote reference number
    reference_number = footnote_reference.reference_id

    # Create a separate paragraph for the footnote text
    footnote_text = doc.add_paragraph(footnote_reference)

    # Set the font size of the footnote text
    for run in footnote_text.runs:
        font = run.font
        font.size = Pt(10)


    heading = "Table"
    doc.add_page_break()
    doc.add_heading(heading, level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a table to the document
    table = doc.add_table(rows=1, cols=len(table_data[0]))
    table.autofit = True

    # Add header row to the table
    header_row = table.rows[0].cells
    for col, header in enumerate(table_data[0]):
        header_row[col].text = str(header)

    # Add data rows to the table
    for row_data in table_data[1:]:
        row_cells = table.add_row().cells
        for col, value in enumerate(row_data):
            row_cells[col].text = str(value)

    merge_cells_in_column(table, column_index=1)
    
    # Apply a built-in table style
    apply_table_style(table)

    # Save the document to the specified output path
    doc.save(output_path)

def merge_cells_in_column(table, column_index):
    column_values = [cell.text for cell in table.columns[column_index].cells]
    merged_cells = set()

    for i, value in enumerate(column_values):
        if value not in merged_cells:
            # Find all cells with the same value in the column
            cells_to_merge = [cell for cell in table.columns[column_index].cells[i:] if cell.text == value]

            # Merge the cells
            if len(cells_to_merge) > 1:
                # Delete the content of all cells except the last one
                for cell in cells_to_merge[:-1]:
                    cell.text = ''
                
                # Update the set of merged cells
                merged_cells.add(value)

# TODO:
# generate chapters in a loop
# chapter should be Chapter class - it contains paragraph, tables, image paths, sources
# loop iterates over list of objects. 
# Each object should representate, "unpack" or print (repr) in chapter in it's specific way
                
def generate_md_chapter(title, heading, paragraph, table, table_name, image, image_description, source):
    md_content = f"# {title}\n\n"

    # Heading
    md_content += f"## {heading}\n\n"

    # Paragraph
    md_content += f"{paragraph}\n\n"
    
    # Table
    md_content += f"Tablica: {table_name} \n"
    md_content += "| Latinski naziv vrste | Hrvatski naziv vrste |\n"
    md_content += "| -------- | -------- |\n"
    for count, row in enumerate(table):
        if row == table[-1]:
            break
        else:
            md_content += f"| {row[0]} | {row[1]} |\n"
    md_content += "\n"

    # Image
    md_content += f"![{image_description}]({image})\n\n"

    # Source as footnote
    md_content += f"Izvor: {source}[^1]\n\n"

    # Footnote
    md_content += "[^1]: Detalji izvora dolaze ovdje u fusnoti \n"

    return md_content


def create_report(title, paragraph, table, image, output):
    chapter = "Natura 2000 područja"
    table_name = "Vrste ptica na predmetnoj lokaciji\n"
    image_description = "Karta obuhvata zahvata u odnosu na Natura2000 područja\n"
    source = "https://www.haop.hr/hr/tematska-podrucja/odrzivo-koristenje-prirodnih-dobara-i-ekoloska-mreza/ekoloska-mreza\n"

    md_content = generate_md_chapter(title, chapter, paragraph, table, table_name, image, image_description, source)

    md_path = "output.md"
    with open(md_path, "w", encoding="utf-8") as md_file:
        md_file.write(md_content)
        
    md_to_docx(md_path, output)


def md_to_docx(md_file, docx_file):
    # Read the Markdown content from the file
    with open(md_file, 'r', encoding='utf-8') as md_file:
        md_content = md_file.read()

    # Convert Markdown to HTML, then html to docx
    html_content = markdown2.markdown(md_content, extras=['tables', 'footnotes'])
    with open("temporary_html_report.html", 'w', encoding='utf-8') as html_file:
        html_file.write(html_content)

    pypandoc.convert_text(html_content, to = "docx", format="html", outputfile=docx_file)


def report_knit(file_to_knit, output, format = 'docx'):
    pypandoc.convert_file(file_to_knit, to = format, outputfile=output)


if __name__ == "__main__":

    table_data = [
    ["ID", "Name", "Value"],
    [1, "Item 1", 10],
    [2, "Item 2", 20],
    [3, "Item 3", 30]
]
    