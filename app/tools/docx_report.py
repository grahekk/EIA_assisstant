from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_existing_section(doc, project, chapter):
    # Access the first paragraph and modify its text and formatting
    first_paragraph = doc.paragraphs[0]
    first_paragraph.text = project.project_title

    map_paragraph = doc.paragraphs[10]
    # Remove the existing paragraph
    map_paragraph.text = ""
    r = map_paragraph.add_run()
    r.add_picture('map_image.jpg', width=Inches(9.0))

    if chapter.table == None:
        pass
    else:
        # Add a table in place of the existing table
        table = doc.tables[0]
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col in range(len(chapter.table_columns)):
            table.cell(0, col).text = chapter.table_columns[col]
        for row_data in chapter.table[0:len(chapter.table)-1]:
            row = table.add_row().cells

        for i, row in enumerate(chapter.table, start=1):
            for col in range(len(chapter.table_columns)):
                try:
                    if i == len(chapter.table):
                        break
                    if str(row[col]) == 'None':
                        continue
                    else:
                        table.cell(i, col).text = str(row[col])
                except TypeError as e:
                    print(f"Skipping {i, row, col} because of {e}")


    source_paragraph = doc.paragraphs[7]
    source_paragraph.text = chapter.table_source
    source_paragraph.style = "Izvor"

    # navigating through paragraphs and document content
    for i, para in enumerate(doc.paragraphs):
        print(i, para.text)

    for j, table in enumerate(doc.tables):
        print(j, table)


def add_chapter(doc, chapter):
    # Add chapter heading
    doc.add_heading(chapter.heading, level=1)

    # Add chapter description
    doc.add_paragraph(chapter.description)

    # Check if chapter has a table
    if chapter.table:
        # Add table heading
        doc.add_paragraph(chapter.table_description)

        # Add table
        try:
            doc.add_paragraph(chapter.table_caption, style = "DE_Tablica CAPTION")
            add_chapters_table(doc, chapter)
            doc.add_paragraph(chapter.table_source, style = 'Izvor')
        except ValueError:
            # Handle if table data is empty or invalid
            doc.add_paragraph("Table data is missing or invalid.")

    try:
        if chapter.image:
            doc.add_picture(chapter.image, width=Inches(5.0))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"Slika: Prikaz gubitaka staništa na tortnom grafu")
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    except Exception as e:
        print(f"Image was skipped because {e}")

def report_from_docx_template(project, report):
    # Open an existing document
    template = r"app\tools\Template.docx"
    doc = Document(template)

    # Iterate over project chapters
    for chapter in project.chapters:
        if chapter.heading == "Zaštićena područja Natura2000":
            # Update existing section if it's the "Natura chapter"
            update_existing_section(doc, project, chapter)
        else:
            # Add a new chapter
            add_chapter(doc, chapter)

    # Save the modified document
    doc.save(report)


def add_chapters_table(doc, chapter):
    table = doc.add_table(rows=len(chapter.table)+1, cols=len(chapter.table_columns))
    table.style = 'DV_tablica'
    table.autofit = True
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)
    for col in range(len(chapter.table_columns)):
        table.cell(0, col).text = chapter.table_columns[col]

    for i, row in enumerate(chapter.table, start=1):
        for col in range(len(chapter.table_columns)):
            try:
                table.cell(i, col).text = str(row[col])
            except TypeError as e:
                print(f"Skipping {i, row, col} because of {e}")
